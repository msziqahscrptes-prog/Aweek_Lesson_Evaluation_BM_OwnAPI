import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import io
import datetime
import time
import json
from google import genai
from google.genai import types

# --- Fungsi Pembantu untuk Gaya Word ---
def set_cell_background(cell, fill_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def format_cell_text(cell, font_size=9.5, bold=False, alignment=None):
    for p in cell.paragraphs:
        if alignment:
            p.alignment = alignment
        for r in p.runs:
            r.font.size = Pt(font_size)
            r.font.name = 'Arial'
            r.font.bold = bold


# --- Enjin Penjana Kandungan Penilaian Resilien Gemini ---
def generate_lesson_evaluation_with_gemini(client, topic, syllabus_code):
    if not topic.strip():
        return "", "", ""

    safe_topic = topic.replace('"', "'").strip()

    # Sandaran sekiranya API gagal (Menggunakan kata 'Pelajar')
    fallback_www = f"Pelajar menghadapi cabaran susun atur asas dan kejelasan konsep teras semasa memetakan '{safe_topic}'."
    fallback_ebw = f"Gunakan contoh semakan sasaran, tugasan penjejakan struktur bersasaran, atau pecahan kolaboratif peraturan '{safe_topic}'."
    fallback_wf = f"Manfaatkan elemen organisasi teras dan struktur logik '{safe_topic}' dalam tugasan pembukaan. Dilaksanakan pada pelajaran seterusnya."

    max_retries = 3
    for attempt in range(max_retries):
        try:
            prompt = f"""
            Analisis topik pelajaran berikut yang diajar di bawah spesifikasi silibus "{syllabus_code or 'Kurikulum Umum'}".
            Topik: "{safe_topic}"

            Sediakan penilaian pedagogi analitikal tersuai khusus untuk topik ini dalam Bahasa Melayu sepenuhnya. 
            PENTING: JANGAN gunakan perkataan 'Murid' sama sekali, gantikan dengan perkataan 'Pelajar'.
            Kembalikan objek JSON dengan kunci tepat berikut:
            - "www": Sesuatu kesukaran, salah faham, atau kesilapan spesifik konteks yang dihadapi oleh pelajar semasa mempelajari topik tepat ini (anggaran 25 patah perkataan).
            - "ebw": Tugasan bilik darjah yang kreatif, praktikal, atau aktiviti semakan untuk menangani isu ini secara langsung (anggaran 25 patah perkataan).
            - "wf": Laluan pemulihan / tindakan susulan. MESTI diakhiri dengan frasa tepat: Dilaksanakan pada pelajaran seterusnya. (anggaran 20 patah perkataan).
            """

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema={
                        "type": "OBJECT",
                        "properties": {
                            "www": {"type": "STRING"},
                            "ebw": {"type": "STRING"},
                            "wf": {"type": "STRING"}
                        },
                        "required": ["www", "ebw", "wf"]
                    },
                    temperature=0.7
                )
            )

            raw_text = response.text.strip()
            if raw_text.startswith("```"):
                raw_text = raw_text.split("```")[1]
                if raw_text.startswith("json"):
                    raw_text = raw_text[4:]

            data = json.loads(raw_text.strip())
            www = data.get("www", "").strip()
            ebw = data.get("ebw", "").strip()
            wf = data.get("wf", "").strip()

            if wf and "Dilaksanakan pada pelajaran seterusnya." not in wf:
                wf = wf.rstrip('.') + ". Dilaksanakan pada pelajaran seterusnya."

            if www and ebw and wf:
                return www, ebw, wf

        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(2.5)
                continue
            break

    return fallback_www, fallback_ebw, fallback_wf


def generate_weekly_conclusion_with_gemini(client, topics, syllabus_code):
    valid_topics = [t for t in topics if t.strip()]
    if not valid_topics:
        return "Tiada sesi pengajaran dijalankan dalam tempoh penilaian minggu ini."
    try:
        prompt = f"Tulis satu kesimpulan ringkasan profesional (sekitar 30 patah perkataan) dalam Bahasa Melayu untuk menilai minggu pengajaran yang merangkumi topik: {', '.join(valid_topics)} di bawah kod silibus {syllabus_code or 'standard'}. Fokus pada pencapaian penting dan pertumbuhan kemahiran pelajar. JANGAN gunakan perkataan 'Murid', gunakan 'Pelajar'."
        return client.models.generate_content(model='gemini-2.5-flash', contents=prompt).text.strip()
    except:
        return f"Objektif utama mingguan subjek yang dipetakan ke parameter silibus berjaya diselesaikan dengan peningkatan kemahiran yang stabil ditunjukkan dalam tugasan pelajar."


# --- Fungsi Membina Dokumen Word ---
def build_word_document(week_no, start_date, end_date, days_data, conclusion):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

    p = doc.add_paragraph()
    p.add_run("PENILAIAN PELAJARAN MINGGUAN UNTUK MINGGU NO : ").font.bold = True
    p.add_run(f"   {week_no}  ").font.underline = True
    p.add_run("\t\t\t\t\t\tTARIKH: ").font.bold = True
    p.add_run(f" {start_date.strftime('%d/%m/%Y')} ").font.underline = True
    p.add_run(" hingga ")
    p.add_run(f" {end_date.strftime('%d/%m/%Y')} ").font.underline = True
    for r in p.runs: r.font.size = Pt(13)

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    col_widths = [Inches(1.2), Inches(1.8), Inches(2.4), Inches(2.4), Inches(2.4)]
    headers = ["TARIKH/HARI", "TOPIK", "APA YANG SALAH (WWW)", "LEBIH BAIK DENGAN (EBW)", "HALA TUJU (WAY FORWARD)"]

    for i, title in enumerate(headers):
        table.rows[0].cells[i].text = title
        format_cell_text(table.rows[0].cells[i], bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, day in enumerate(days_data):
        row_cells = table.rows[idx + 1].cells
        for col_idx, text in enumerate([day["date_str"], day["topic"], day["www"], day["ebw"], day["wf"]]):
            row_cells[col_idx].text = text

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell)
            format_cell_text(cell)

    doc.add_paragraph(
        "\nPenilaian keseluruhan adalah pilihan. Sebagai contoh, jika anda ingin merekodkan 'apa yang berjalan lancar' atau butiran lain pengajaran anda, anda boleh memasukkannya di penghujung, seperti berikut.").runs[
        0].font.size = Pt(9.5)

    bottom_table = doc.add_table(rows=2, cols=3)
    bottom_table.style = 'Table Grid'
    bottom_widths = [Inches(1.8), Inches(5.8), Inches(2.4)]

    bottom_table.rows[0].cells[0].text = "APA YANG BERJALAN LANCAR\n\n(Kesimpulan)"
    bottom_table.rows[0].cells[1].text = conclusion
    bottom_table.rows[1].cells[0].text = "CATATAN / REMARK"
    bottom_table.rows[1].cells[2].text = "\n\nNama/Tandatangan"

    set_cell_background(bottom_table.rows[0].cells[0], "66FFCC")
    set_cell_background(bottom_table.rows[1].cells[0], "66FFCC")

    for row in bottom_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = bottom_widths[i]
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            align = WD_ALIGN_PARAGRAPH.RIGHT if i == 2 and row == bottom_table.rows[1] else None
            format_cell_text(cell, bold=(i == 0), alignment=align)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- Antara Muka Aplikasi Web Streamlit ---
st.set_page_config(layout="wide", page_title="Portal Penilaian Pelajaran AI")
st.title("📋 Portal Automasi Penilaian Rancangan Mengajar (BM)")

# --- INPUT KUNCI API ---
user_api_key = st.text_input(
    "🔑 Masukkan Kunci API Gemini Anda:", 
    type="password", 
    help="Dapatkan kunci API peribadi anda dari Google AI Studio menggunakan akaun Gmail anda."
)

st.subheader("Panel Konfigurasi")
c1, c2, c3 = st.columns(3)
week_input = c1.text_input("MINGGU NO:", value="01")
start_dt = c2.date_input("Tarikh Mula Isnin:", datetime.date.today())
syllabus_input = c3.text_input("KOD SILIBUS / SPESIFIKASI:", value="9696")

day_offsets = {"Isnin": 0, "Selasa": 1, "Rabu": 2, "Khamis": 3, "Sabtu": 5}
end_dt = start_dt + datetime.timedelta(days=5)
st.info(
    f"Sesi Penilaian Sasaran: **{start_dt.strftime('%d/%m/%Y')}** hingga **{end_dt.strftime('%d/%m/%Y')}** | Kerangka Rujukan: **{syllabus_input if syllabus_input else 'Kurikulum Umum'}**")

st.markdown("---")
st.write("##### Input Topik Pelajaran")
topics_collected = {}
col_l, col_r = st.columns(2)

for idx, (day_name, offset) in enumerate(day_offsets.items()):
    target_pane = col_l if idx < 3 else col_r
    calc_date = start_dt + datetime.timedelta(days=offset)
    date_str = calc_date.strftime("%d/%m/%Y")
    user_topic = target_pane.text_input(f"Topik {day_name} ({date_str}):", key=f"inp_{day_name}")
    topics_collected[day_name] = {"date_str": f"{date_str}\n({day_name})", "topic": user_topic}

if st.button("JANAKAN LAPORAN PENILAIAN", type="primary"):
    if not user_api_key:
        st.error("❌ Ralat Konfigurasi Kunci! Sila masukkan kunci API Gemini peribadi anda di bahagian atas sebelum menjana laporan.")
    else:
        try:
            # Inisialisasi genai.Client dengan kunci input pengguna
            client_instance = genai.Client(api_key=user_api_key)
            
            report_rows, raw_topics_list = [], []
            status_box = st.empty()

            processed_count = 0
            for day_name, d_info in topics_collected.items():
                topic_text = d_info["topic"]
                raw_topics_list.append(topic_text)

                if topic_text.strip():
                    if processed_count > 0:
                        for remaining in range(6, 0, -1):
                            status_box.info(f"⏳ Menstabilkan saluran sistem ({remaining}s)... Menyediakan: **{day_name}**")
                            time.sleep(1)

                    with st.spinner(f"🚀 Gemini sedang menganalisis: **{day_name}** ({topic_text})..."):
                        www_out, ebw_out, wf_out = generate_lesson_evaluation_with_gemini(client_instance, topic_text,
                                                                                          syllabus_input)
                    processed_count += 1
                else:
                    www_out, ebw_out, wf_out = "", "", ""

                report_rows.append(
                    {"date_str": d_info["date_str"], "topic": topic_text, "www": www_out, "ebw": ebw_out, "wf": wf_out})

            with st.spinner("✍️ Menulis kesimpulan kotak ringkasan akhir..."):
                final_conclusion = generate_weekly_conclusion_with_gemini(client_instance, raw_topics_list, syllabus_input)
            status_box.empty()

            st.session_state.update(
                {'data_processed': True, 'report_rows': report_rows, 'final_conclusion': final_conclusion})
                
        except Exception as api_err:
            st.error(f"❌ Ralat Sambungan atau Pengesahan API: {str(api_err)}")

if st.session_state.get('data_processed'):
    st.markdown("---")
    st.subheader("👁️ Pratonton Dokumen Atas Skrin")
    st.table(st.session_state['report_rows'])
    st.write("**Kotak Ringkasan Di Janakan:**", st.session_state['final_conclusion'])

    word_file = build_word_document(week_input, start_dt, end_dt, st.session_state['report_rows'],
                                    st.session_state['final_conclusion'])
    st.download_button(label="📥 Muat Turun Dokumen Word Rasmi (.docx)", data=word_file,
                       file_name=f"Penilaian_Pelajaran_Minggu_{week_input}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
